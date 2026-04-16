import fitz  # PyMuPDF
from docx import Document
import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text)

        pdf_document.close()
        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("PDF appears to be image-based or empty. Please upload a text-based PDF.")

        logger.info(f"Extracted {len(full_text)} characters from PDF")
        return full_text.strip()

    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("DOCX file appears to be empty.")

        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text.strip()

    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def parse_file(file_bytes: bytes, filename: str) -> str:
    """Parse file based on extension and return extracted text."""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx or .pdf")
    else:
        raise ValueError(f"Unsupported file format: {filename}. Please upload a PDF or DOCX file.")